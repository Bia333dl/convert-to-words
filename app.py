import streamlit as st
from google import genai
from PIL import Image
from docx import Document
import img2pdf
from pdf2image import convert_from_path
import os
import tempfile
import io

# ===================== TABLE FUNCTION (Defined First) =====================
def add_markdown_table_to_doc(doc, table_lines):
    """Convert Markdown table to real Word table"""
    if len(table_lines) < 2:
        return
    clean_lines = [line for line in table_lines if '---' not in line]
    if len(clean_lines) < 1:
        return
    header = [cell.strip() for cell in clean_lines[0].split('|')[1:-1]]
    rows = []
    for line in clean_lines[2:]:
        cells = [cell.strip() for cell in line.split('|')[1:-1]]
        if cells:
            rows.append(cells)
    cols = len(header) if header else (len(rows[0]) if rows else 1)
    table = doc.add_table(rows=1 + len(rows), cols=cols)
    table.style = 'Table Grid'
    if header:
        for j, text in enumerate(header):
            table.cell(0, j).text = text
    for i, row in enumerate(rows):
        for j, text in enumerate(row):
            if j < cols:
                table.cell(i + 1, j).text = text

# ===================== LANGUAGE SETUP =====================
if "language" not in st.session_state:
    st.session_state.language = "vi"

lang = st.session_state.language

texts = {
    "en": {
        "title": "📄 AI Document Digitizer Pro",
        "subtitle": "Upload scanned PDF or images → Get a well-formatted Word document with tables preserved",
        "api_label": "Enter your Gemini API Key",
        "step1": "Step 1: Combine Images into One PDF (Optional)",
        "upload_images": "Upload multiple images",
        "create_pdf": "Create PDF from Images",
        "step2": "Step 2: Digitize Document with AI",
        "upload_file": "Upload PDF or Image",
        "start_btn": "🚀 Start AI Digitization",
        "success": "✅ Done!",
        "download_word": "📥 Download Structured Word File",
        "download_pdf": "📥 Download Combined PDF"
    },
    "vi": {
        "title": "📄 AI Chuyển Đổi Tài Liệu Pro",
        "subtitle": "Chuyển PDF scan hoặc ảnh → File Word có định dạng đẹp, giữ bảng biểu",
        "api_label": "Nhập Gemini API Key của bạn",
        "step1": "Bước 1: Ghép nhiều ảnh thành 1 PDF (Tùy chọn)",
        "upload_images": "Chọn nhiều ảnh để ghép",
        "create_pdf": "🔨 Tạo PDF từ Ảnh",
        "step2": "Bước 2: Chuyển Đổi Tài Liệu Bằng AI",
        "upload_file": "Tải lên PDF hoặc Ảnh",
        "start_btn": "🚀 Bắt Đầu Chuyển Đổi Bằng AI",
        "success": "✅ Hoàn thành!",
        "download_word": "📥 Tải File Word Đã Định Dạng",
        "download_pdf": "📥 Tải PDF Đã Ghép"
    }
}

t = texts[lang]

st.set_page_config(page_title=t["title"], layout="centered")

# Language Switcher
col1, col2 = st.columns([4, 1])
with col1:
    st.title(t["title"])
with col2:
    if st.button("🇻🇳 VN" if lang == "en" else "🇬🇧 EN"):
        st.session_state.language = "en" if lang == "vi" else "vi"
        st.rerun()

st.markdown(f"**{t['subtitle']}**")

# API Key Guide
with st.expander("📋 How to Get Gemini API Key / Hướng dẫn lấy API Key", expanded=False):
    col1, col2 = st.columns(2)
    with col1:
        st.markdown("**🇬🇧 English**")
        st.markdown("1. Go to [Google AI Studio](https://aistudio.google.com/app/apikey)\n2. Sign in\n3. Click Create API key\n4. Copy key")
    with col2:
        st.markdown("**🇻🇳 Tiếng Việt**")
        st.markdown("1. Truy cập [Google AI Studio](https://aistudio.google.com/app/apikey)\n2. Đăng nhập\n3. Nhấn Create API key\n4. Sao chép key")

api_key = st.text_input(t["api_label"], type="password")

# Step 1: Combine Images
st.subheader(t["step1"])
image_files = st.file_uploader(t["upload_images"], type=["jpg","jpeg","png","webp"], accept_multiple_files=True, key="images")

if st.button(t["create_pdf"]) and image_files:
    with st.spinner("Đang tạo PDF..." if lang == "vi" else "Creating PDF..."):
        try:
            with tempfile.TemporaryDirectory() as tmpdir:
                paths = []
                for img in image_files:
                    path = os.path.join(tmpdir, img.name)
                    with open(path, "wb") as f:
                        f.write(img.getbuffer())
                    paths.append(path)
                pdf_bytes = img2pdf.convert(paths)
                st.download_button(
                    label=t["download_pdf"],
                    data=pdf_bytes,
                    file_name="tai_lieu_da_ghep.pdf" if lang == "vi" else "combined.pdf",
                    mime="application/pdf"
                )
                st.success("✅ Đã tạo PDF thành công!" if lang == "vi" else "✅ PDF created successfully!")
        except Exception as e:
            st.error(str(e))

# Step 2: Main Tool
st.subheader(t["step2"])
uploaded_file = st.file_uploader(t["upload_file"], type=["pdf","jpg","jpeg","png","webp"], key="main")

if st.button(t["start_btn"], type="primary") and uploaded_file and api_key:
    with st.spinner("Đang xử lý bằng Gemini AI..." if lang == "vi" else "Processing with Gemini AI..."):
        try:
            client = genai.Client(api_key=api_key)
            
            suffix = os.path.splitext(uploaded_file.name)[1].lower()
            with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                tmp.write(uploaded_file.getbuffer())
                file_path = tmp.name

            if suffix == '.pdf':
                try:
                    images_to_process = convert_from_path(file_path)
                except:
                    poppler_path = r"E:\python project\convert to word\poppler-24.08.0\Library\bin"
                    images_to_process = convert_from_path(file_path, poppler_path=poppler_path)
            else:
                images_to_process = [Image.open(file_path)]

            prompt = "Analyze the document carefully. Preserve original layout. Transcribe all text accurately. Convert tables into proper Markdown table format."

            contents = [prompt] + images_to_process
            response = client.models.generate_content(model='gemini-2.5-flash', contents=contents)
            
            extracted_text = response.text if response and response.text else "No text extracted."

            doc = Document()
            lines = extracted_text.splitlines()
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                if line.startswith('#'):
                    level = min(line.count('#'), 3)
                    text = line.lstrip('# ').strip()
                    doc.add_heading(text, level=level)
                elif line.startswith('|') and len(line) > 10:
                    table_lines = []
                    while i < len(lines) and lines[i].strip().startswith('|'):
                        table_lines.append(lines[i].strip())
                        i += 1
                    add_markdown_table_to_doc(doc, table_lines)
                    continue
                elif line:
                    doc.add_paragraph(line)
                i += 1

            bio = io.BytesIO()
            doc.save(bio)
            bio.seek(0)

            st.success(t["success"])
            st.download_button(
                label=t["download_word"],
                data=bio.getvalue(),
                file_name=f"{os.path.splitext(uploaded_file.name)[0]}_structured.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            os.unlink(file_path)

        except Exception as e:
            st.error(f"❌ Lỗi: {str(e)}" if lang == "vi" else f"❌ Error: {str(e)}")

st.caption("Công cụ chuyển đổi tài liệu scan thông minh | Powered by Gemini AI")
